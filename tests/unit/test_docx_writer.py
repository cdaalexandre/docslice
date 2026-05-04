"""Tests - write_txt_as_docx with real files.

python-docx is a pure-library with no external I/O, so we test the
real function against tmp_path. No Fakes needed (Percival & Gregory,
Cap. 3.c.ii: prefer real over mock when there is no boundary).
"""

from __future__ import annotations

from typing import TYPE_CHECKING

import pytest
from docx import Document

from docslice.adapters.docx_writer import write_txt_as_docx

if TYPE_CHECKING:
    from pathlib import Path


class TestWriteTxtAsDocx:
    """Tests for the TXT -> .docx adapter."""

    def test_creates_docx_file(self, tmp_path: Path) -> None:
        txt = tmp_path / "input.txt"
        txt.write_text("Hello world.", encoding="utf-8")
        docx = tmp_path / "out.docx"

        result = write_txt_as_docx(txt, docx)

        assert result == docx
        assert docx.exists()
        assert docx.stat().st_size > 0

    def test_one_paragraph_per_blank_line_block(self, tmp_path: Path) -> None:
        txt = tmp_path / "input.txt"
        txt.write_text(
            "First paragraph here.\n\nSecond paragraph here.\n\nThird one.",
            encoding="utf-8",
        )
        docx = tmp_path / "out.docx"

        write_txt_as_docx(txt, docx)

        document = Document(str(docx))
        bodies = [p.text for p in document.paragraphs]
        assert bodies == [
            "First paragraph here.",
            "Second paragraph here.",
            "Third one.",
        ]

    def test_skips_empty_blocks(self, tmp_path: Path) -> None:
        txt = tmp_path / "input.txt"
        txt.write_text(
            "Real paragraph.\n\n\n\n   \n\nAnother real one.",
            encoding="utf-8",
        )
        docx = tmp_path / "out.docx"

        write_txt_as_docx(txt, docx)

        document = Document(str(docx))
        bodies = [p.text for p in document.paragraphs]
        assert bodies == ["Real paragraph.", "Another real one."]

    def test_preserves_unicode(self, tmp_path: Path) -> None:
        txt = tmp_path / "input.txt"
        txt.write_text("Cafe com canela. Sao Paulo, SP. Acao.", encoding="utf-8")
        docx = tmp_path / "out.docx"

        write_txt_as_docx(txt, docx)

        document = Document(str(docx))
        assert document.paragraphs[0].text == "Cafe com canela. Sao Paulo, SP. Acao."

    def test_creates_parent_directory(self, tmp_path: Path) -> None:
        txt = tmp_path / "input.txt"
        txt.write_text("Content.", encoding="utf-8")
        docx = tmp_path / "deep" / "nested" / "out.docx"

        write_txt_as_docx(txt, docx)

        assert docx.exists()

    def test_missing_input_raises(self, tmp_path: Path) -> None:
        txt = tmp_path / "missing.txt"
        docx = tmp_path / "out.docx"

        with pytest.raises(FileNotFoundError):
            write_txt_as_docx(txt, docx)
